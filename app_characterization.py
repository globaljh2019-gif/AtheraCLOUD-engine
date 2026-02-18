import streamlit as st
import pandas as pd
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import io
from datetime import datetime

# ==========================================
# 0. Notion Master Blueprint 기반 지식 베이스
# ==========================================
def get_notion_master_db(lang_code):
    if lang_code == "KR":
        return [
            {"Category": "1. 구조적 특성", "Attribute": "1차 구조 (아미노산 서열)", "Method": "Peptide Mapping (LC-MS/MS)", "Tier": "Tier 1", "Rationale": "아미노산 서열 일치성 및 PTM 확인 필수", "Dev_Strategy": "Trypsin 소화 효율 최적화 및 Coverage 95% 이상 확보 전략."},
            {"Category": "1. 구조적 특성", "Attribute": "당쇄 프로파일 (N-Glycan)", "Method": "HILIC-FLD / MS", "Tier": "Tier 1", "Rationale": "면역원성 및 이펙터 기능(ADCC) 영향 분석", "Dev_Strategy": "2-AB 라벨링 효율 및 주요 당쇄(G0F, G1F 등) 분리능 최적화."},
            {"Category": "2. 물리화학적 성질", "Attribute": "전하 변이체 (Charge Variants)", "Method": "CEX-HPLC / cIEF", "Tier": "Tier 1", "Rationale": "단백질 안정성 및 불순물 프로파일 확인", "Dev_Strategy": "pH Gradient를 이용한 Acidic/Basic 변이체 분리능 극대화."},
            {"Category": "2. 물리화학적 성질", "Attribute": "크기 변이체 (응집체)", "Method": "SEC-HPLC", "Tier": "Tier 1", "Rationale": "단백질 응집에 따른 안전성 위험 관리", "Dev_Strategy": "비특이적 결합 방지를 위한 이동상 염 농도 및 유속 최적화."},
            {"Category": "3. 생물학적 활성", "Attribute": "결합 역가 (Binding Affinity)", "Method": "SPR (Biacore) / ELISA", "Tier": "Tier 1", "Rationale": "항원-항체 결합력(KD) 및 특이성 입증", "Dev_Strategy": "Chip 표면 고정화 농도 최적화 및 Kinetics 분석 정밀도 확보."},
        ]
    else:
        return [
            {"Category": "1. Structural", "Attribute": "Primary Structure", "Method": "Peptide Mapping (LC-MS/MS)", "Tier": "Tier 1", "Rationale": "Sequence confirmation and PTM site mapping", "Dev_Strategy": "Optimize digestion and target >95% sequence coverage."},
            {"Category": "1. Structural", "Attribute": "Glycan Profile (N-linked)", "Method": "HILIC-FLD / MS", "Tier": "Tier 1", "Rationale": "Impact on immunogenicity and ADCC activity", "Dev_Strategy": "Maximize labeling efficiency and resolve major glycoforms."},
            {"Category": "2. Physicochemical", "Attribute": "Charge Variants", "Method": "CEX-HPLC / cIEF", "Tier": "Tier 1", "Rationale": "Assessment of stability and variant profile", "Dev_Strategy": "Optimize pH gradient for acidic/basic peak resolution."},
            {"Category": "2. Physicochemical", "Attribute": "Size Variants (Aggregates)", "Method": "SEC-HPLC", "Tier": "Tier 1", "Rationale": "Safety risk management for protein aggregation", "Dev_Strategy": "Screen mobile phase salt concentration to prevent non-specific binding."},
            {"Category": "3. Biological", "Attribute": "Binding Affinity", "Method": "SPR (Biacore) / ELISA", "Tier": "Tier 1", "Rationale": "Demonstrate antigen-antibody binding (KD)", "Dev_Strategy": "Optimize ligand density and ensure kinetic data quality."},
        ]

# ==========================================
# 1. 지식 베이스 (Database - Dual Language)
# ==========================================
def get_method_database(modality, lang):
    """
    모달리티별 시험 항목 DB (국문/영문 스위칭)
    """
    if modality == "Monoclonal Antibody (mAb)":
        if lang == "KR":
            # [국문 데이터]
            data = [
                {"Category": "1. 구조적 특성", "Attribute": "1차 구조 (아미노산 서열)", "Method": "Peptide Mapping (LC-MS/MS)", "Tier": "필수 (Tier 1)", "Dev_Strategy": "Trypsin 소화 효율 최적화 (4시간 vs Overnight). Sequence Coverage 95% 이상 목표."},
                {"Category": "1. 구조적 특성", "Attribute": "고차 구조 (2차/3차)", "Method": "CD (Far/Near UV) & DSC", "Tier": "심화 (Tier 2)", "Dev_Strategy": "Buffer 간섭 최소화 및 Reference와의 스펙트럼 중첩성(Similarity) 비교."},
                {"Category": "1. 구조적 특성", "Attribute": "이황화 결합", "Method": "Non-reduced / Reduced Peptide Mapping", "Tier": "필수 (Tier 1)", "Dev_Strategy": "Free Thiol 측정 병행. Scrambled disulfide bond 유무 확인."},
                {"Category": "1. 구조적 특성", "Attribute": "당쇄 프로파일 (N-Glycan)", "Method": "HILIC-FLD / MS", "Tier": "필수 (Tier 1)", "Dev_Strategy": "주요 당쇄(G0F, G1F 등) 정량 및 면역원성 당쇄(Man5, G0) 모니터링."},
                {"Category": "2. 물리화학적 성질", "Attribute": "전하 변이체", "Method": "CEX-HPLC (Salt/pH Gradient)", "Tier": "필수 (Tier 1)", "Dev_Strategy": "Acidic/Basic peak 분리능 확보. 등전점(pI) 확인."},
                {"Category": "2. 물리화학적 성질", "Attribute": "크기 변이체 (응집체)", "Method": "SEC-HPLC", "Tier": "필수 (Tier 1)", "Dev_Strategy": "비특이적 결합 방지(염 농도 조절). HMW/Monomer 분리능 확인."},
                {"Category": "2. 물리화학적 성질", "Attribute": "크기 변이체 (분해물)", "Method": "CE-SDS (Non-reduced)", "Tier": "필수 (Tier 1)", "Dev_Strategy": "샘플 전처리 온도/시간 최적화로 인위적 분해 방지."},
                {"Category": "3. 생물학적 활성", "Attribute": "결합 활성 (Binding)", "Method": "ELISA / SPR", "Tier": "필수 (Tier 1)", "Dev_Strategy": "항원 코팅 농도 최적화 및 평행성(Parallelism) 입증."},
                {"Category": "3. 생물학적 활성", "Attribute": "작용 기전 역가 (Potency)", "Method": "Cell-based Assay", "Tier": "필수 (Tier 1)", "Dev_Strategy": "세포주 민감도 확인 및 4-PL 커브 피팅 적합성 평가."},
                {"Category": "4. 불순물", "Attribute": "공정 유래 불순물", "Method": "HCP ELISA & qPCR", "Tier": "필수 (Tier 1)", "Dev_Strategy": "공정 특이적 키트 선정 및 DNA 추출 효율 확인."},
            ]
        else:
            # [English Data]
            data = [
                {"Category": "1. Structure", "Attribute": "Primary Structure", "Method": "Peptide Mapping (LC-MS/MS)", "Tier": "Tier 1", "Dev_Strategy": "Optimize digestion efficiency (4h vs overnight). Target >95% sequence coverage."},
                {"Category": "1. Structure", "Attribute": "Higher Order Structure", "Method": "CD (Far/Near UV) & DSC", "Tier": "Tier 2", "Dev_Strategy": "Minimize buffer interference. Compare spectral similarity with reference standard."},
                {"Category": "1. Structure", "Attribute": "Disulfide Bond", "Method": "Non-reduced / Reduced Mapping", "Tier": "Tier 1", "Dev_Strategy": "Check free thiols (Ellman's). Confirm absence of scrambled bonds."},
                {"Category": "1. Structure", "Attribute": "Glycan Profile", "Method": "HILIC-FLD / MS", "Tier": "Tier 1", "Dev_Strategy": "Quantify major glycans (G0F, G1F) and monitor immunogenic species (Man5)."},
                {"Category": "2. Physicochemical", "Attribute": "Charge Variants", "Method": "CEX-HPLC", "Tier": "Tier 1", "Dev_Strategy": "Ensure resolution of Acidic/Basic peaks. Verify pI consistency."},
                {"Category": "2. Physicochemical", "Attribute": "Size Variants (Aggregates)", "Method": "SEC-HPLC", "Tier": "Tier 1", "Dev_Strategy": "Control salt conc. to prevent non-specific binding. Check resolution."},
                {"Category": "2. Physicochemical", "Attribute": "Size Variants (Fragments)", "Method": "CE-SDS (Non-reduced)", "Tier": "Tier 1", "Dev_Strategy": "Optimize sample prep temp/time to minimize artificial degradation."},
                {"Category": "3. Biological Activity", "Attribute": "Binding Activity", "Method": "ELISA / SPR", "Tier": "Tier 1", "Dev_Strategy": "Optimize coating concentration. Demonstrate parallelism."},
                {"Category": "3. Biological Activity", "Attribute": "Potency (MoA)", "Method": "Cell-based Assay", "Tier": "Tier 1", "Dev_Strategy": "Check cell line sensitivity. Evaluate 4-PL curve fit suitability."},
                {"Category": "4. Impurities", "Attribute": "Process Impurities", "Method": "HCP ELISA & qPCR", "Tier": "Tier 1", "Dev_Strategy": "Select process-specific kit. Verify DNA recovery efficiency."},
            ]
        return pd.DataFrame(data)
    else:
        return pd.DataFrame()

# ==========================================
# 2. 문서 생성 엔진 (Report Generator - Dual)
# ==========================================

def set_cell_background(cell, color_hex):
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    cell._element.get_or_add_tcPr().append(shd)

def generate_report(product_name, modality, phase, selected_methods, lang):
    doc = Document()
    style = doc.styles['Normal']
    
    # 언어별 텍스트 설정
    if lang == "KR":
        font_name = 'Malgun Gothic'
        title_text = f'{product_name} 특성분석 종합 계획서'
        labels = {"Prod": "제품명", "Mod": "모달리티", "Phase": "단계", "Date": "날짜"}
        headers = ['구분', '품질 속성', '시험 방법', '중요도']
        sec1_title = '1. 특성분석 종합 계획'
        sec1_desc = f"본 문서는 {product_name}의 {phase} 승인을 위한 시험 항목을 정의합니다."
        sec2_title = '2. 시험법 선정 근거'
        sec2_desc = "ICH Q6B 가이드라인 및 CQA 평가에 기반하여 선정됨."
        sec3_title = '3. 개발 전략'
        sec3_desc = "시험법 최적화를 위한 전략:"
        sign_text = "작성자: ___________________  승인자: ___________________"
    else:
        font_name = 'Arial'
        title_text = f'{product_name} Characterization Plan'
        labels = {"Prod": "Product", "Mod": "Modality", "Phase": "Phase", "Date": "Date"}
        headers = ['Category', 'Attribute', 'Method', 'Tier']
        sec1_title = '1. Comprehensive Characterization Plan'
        sec1_desc = f"This document defines the characterization items for {product_name} ({phase})."
        sec2_title = '2. Rationale for Selection'
        sec2_desc = "Selected based on ICH Q6B guidelines and CQA assessment."
        sec3_title = '3. Development Strategy'
        sec3_desc = "Strategies for method optimization:"
        sign_text = "Prepared by: ___________________  Approved by: ___________________"

    # 폰트 적용
    style.font.name = font_name
    if lang == "KR":
        style._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    style.font.size = Pt(10)

    # 타이틀
    title = doc.add_heading(title_text, 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("")
    
    # 정보 테이블
    table_info = doc.add_table(rows=3, cols=2)
    table_info.style = 'Table Grid'
    info_rows = [
        (labels["Prod"], product_name),
        (labels["Mod"], modality),
        (labels["Phase"], phase)
    ]
    for i, (l, v) in enumerate(info_rows):
        table_info.rows[i].cells[0].text = l
        table_info.rows[i].cells[1].text = v
        set_cell_background(table_info.rows[i].cells[0], 'F2F2F2')

    doc.add_paragraph("")

    # Section 1
    doc.add_heading(sec1_title, level=1)
    doc.add_paragraph(sec1_desc)
    
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        set_cell_background(cell, 'E7E6E6')

    for idx, row in selected_methods.iterrows():
        cells = table.add_row().cells
        cells[0].text = str(row['Category'])
        cells[1].text = str(row['Attribute'])
        cells[2].text = str(row['Method'])
        cells[3].text = str(row['Tier'])

    doc.add_paragraph("")

    # Section 2 & 3
    doc.add_heading(sec2_title, level=1)
    doc.add_paragraph(sec2_desc)
    
    doc.add_heading(sec3_title, level=1)
    doc.add_paragraph(sec3_desc)
    for idx, row in selected_methods.iterrows():
        p = doc.add_paragraph(style="List Bullet")
        runner = p.add_run(f"[{row['Method']}] : {row['Dev_Strategy']}")
        runner.bold = False

    doc.add_paragraph("-" * 70)
    doc.add_paragraph(sign_text)

    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

# ==========================================
# 3. 메인 UI (Streamlit - Dual)
# ==========================================
def main():
    st.set_page_config(page_title="AtheraCLOUD Characterization", layout="wide")
    
    with st.sidebar:
        st.title("🧬 AtheraCLOUD")
        
        # [핵심 기능] 언어 선택 스위치
        lang = st.radio("Language / 언어", ["Korean (국문)", "English (영문)"])
        lang_code = "KR" if "Korean" in lang else "EN"

        st.markdown("---")
        
        # 사이드바 라벨도 언어에 따라 변경
        if lang_code == "KR":
            st.subheader("프로젝트 설정")
            modality = st.selectbox("모달리티", ["Monoclonal Antibody (mAb)", "ADC (준비중)"])
            product_name = st.text_input("제품명", "Athera-mAb-001")
            phase = st.selectbox("개발 단계", ["비임상", "임상 1상", "임상 3상", "BLA"])
        else:
            st.subheader("Project Settings")
            modality = st.selectbox("Modality", ["Monoclonal Antibody (mAb)", "ADC (Coming Soon)"])
            product_name = st.text_input("Product Name", "Athera-mAb-001")
            phase = st.selectbox("Phase", ["Pre-clinical", "Phase 1", "Phase 3", "BLA"])

    # 메인 타이틀
    if lang_code == "KR":
        st.markdown(f"## 🧪 {modality} 특성분석 엔진")
        st.markdown("**진행 순서:** 1.항목선정 ➔ 2.개발전략 ➔ 3.리포트")
        tab_names = ["1️⃣ 항목 선정 (Decision)", "2️⃣ 개발 전략 (Guide)", "3️⃣ 리포트 (Report)"]
    else:
        st.markdown(f"## 🧪 {modality} Characterization Engine")
        st.markdown("**Process:** 1.Decision ➔ 2.Strategy ➔ 3.Report")
        tab_names = ["1️⃣ Decision", "2️⃣ Strategy", "3️⃣ Report"]

    # 데이터 로드 (언어 선택 적용)
    df_db = get_method_database("Monoclonal Antibody (mAb)", lang_code)
    
    tab1, tab2, tab3 = st.tabs(tab_names)

    # --- Tab 1: Decision ---
    with tab1:
        if lang_code == "KR":
            st.subheader("시험 항목 선정")
            st.markdown("프로젝트에 필요한 분석 항목을 체크하세요.")
            col_config = {
                "Include": st.column_config.CheckboxColumn("선택"),
                "Category": st.column_config.TextColumn("분류"),
                "Attribute": st.column_config.TextColumn("품질 속성"),
                "Method": st.column_config.TextColumn("시험법"),
                "Tier": st.column_config.TextColumn("중요도")
            }
        else:
            st.subheader("Method Selection")
            st.markdown("Select analysis items for your project.")
            col_config = {
                "Include": st.column_config.CheckboxColumn("Select"),
                "Category": st.column_config.TextColumn("Category"),
                "Attribute": st.column_config.TextColumn("Attribute"),
                "Method": st.column_config.TextColumn("Method"),
                "Tier": st.column_config.TextColumn("Tier")
            }

        df_db['Include'] = True
        edited_df = st.data_editor(
            df_db[['Include', 'Category', 'Attribute', 'Method', 'Tier']],
            column_config=col_config,
            use_container_width=True,
            hide_index=True
        )
        selected_rows = edited_df[edited_df['Include'] == True]

    # --- Tab 2: Strategy ---
    with tab2:
        if lang_code == "KR":
            st.subheader("시험법 개발 전략")
        else:
            st.subheader("Development Strategy")

        if len(selected_rows) > 0:
            final_selection = pd.merge(selected_rows, df_db, on=['Category', 'Attribute', 'Method', 'Tier'], how='left')
            for index, row in final_selection.iterrows():
                strategy = row.get('Dev_Strategy_y', row.get('Dev_Strategy', ''))
                with st.expander(f"📌 {row['Attribute']} - {row['Method']}"):
                    st.info(strategy)
        else:
            st.warning("Please select items in Tab 1.")

    # --- Tab 3: Report ---
    with tab3:
        if lang_code == "KR":
            st.subheader("종합계획서 생성")
            btn_label = "📄 국문 리포트 다운로드 (.docx)"
            file_suffix = "_KR.docx"
        else:
            st.subheader("Generate Report")
            btn_label = "📄 Download English Report (.docx)"
            file_suffix = "_EN.docx"

        if len(selected_rows) > 0:
            final_selection = pd.merge(selected_rows, df_db, on=['Category', 'Attribute', 'Method', 'Tier'], how='left')
            if 'Dev_Strategy_y' in final_selection.columns:
                 final_selection['Dev_Strategy'] = final_selection['Dev_Strategy_y']

            doc_file = generate_report(product_name, modality, phase, final_selection, lang_code)
            
            st.dataframe(selected_rows[['Category', 'Attribute', 'Method']], use_container_width=True, hide_index=True)
            st.download_button(
                label=btn_label,
                data=doc_file,
                file_name=f"{product_name}_Characterization_Plan{file_suffix}",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

if __name__ == "__main__":
    main()