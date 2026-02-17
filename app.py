import streamlit as st
import pandas as pd
import requests
import io
import xlsxwriter
import random
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---------------------------------------------------------
# 0. 페이지 설정
# ---------------------------------------------------------
st.set_page_config(page_title="AtheraCLOUD Validation Suite", layout="wide")

# ---------------------------------------------------------
# 1. 설정 및 데이터 로딩
# ---------------------------------------------------------
try:
    NOTION_API_KEY = st.secrets["NOTION_API_KEY"]
    CRITERIA_DB_ID = st.secrets["CRITERIA_DB_ID"]
    STRATEGY_DB_ID = st.secrets["STRATEGY_DB_ID"]
    PARAM_DB_ID = st.secrets.get("PARAM_DB_ID", "") 
except:
    NOTION_API_KEY = ""
    CRITERIA_DB_ID = ""
    STRATEGY_DB_ID = ""
    PARAM_DB_ID = ""

headers = {"Authorization": "Bearer " + NOTION_API_KEY, "Content-Type": "application/json", "Notion-Version": "2022-06-28"}

@st.cache_data
def get_criteria_map():
    if not CRITERIA_DB_ID: return {}
    url = f"https://api.notion.com/v1/databases/{CRITERIA_DB_ID}/query"
    res = requests.post(url, headers=headers)
    criteria_map = {}
    if res.status_code == 200:
        for p in res.json().get("results", []):
            try:
                props = p["properties"]
                cat = props["Test_Category"]["title"][0]["text"]["content"] if props["Test_Category"]["title"] else "Unknown"
                req = [i["name"] for i in props["Required_Items"]["multi_select"]]
                criteria_map[p["id"]] = {"Category": cat, "Required_Items": req}
            except: continue
    return criteria_map

def get_strategy_list(criteria_map):
    if not STRATEGY_DB_ID: return pd.DataFrame()
    url = f"https://api.notion.com/v1/databases/{STRATEGY_DB_ID}/query"
    res = requests.post(url, headers=headers)
    data = []
    if res.status_code == 200:
        for p in res.json().get("results", []):
            try:
                props = p["properties"]
                mod = props["Modality"]["select"]["name"] if props["Modality"]["select"] else ""
                ph = props["Phase"]["select"]["name"] if props["Phase"]["select"] else ""
                met = props["Method Name"]["rich_text"][0]["text"]["content"] if props["Method Name"]["rich_text"] else ""
                rel = props["Test Category"]["relation"]
                cat, items = ("Unknown", [])
                if rel and rel[0]["id"] in criteria_map:
                    cat = criteria_map[rel[0]["id"]]["Category"]
                    items = criteria_map[rel[0]["id"]]["Required_Items"]
                data.append({"Modality": mod, "Phase": ph, "Method": met, "Category": cat, "Required_Items": items})
            except: continue
    return pd.DataFrame(data)

def get_method_params(method_name):
    if not PARAM_DB_ID: return {}
    url = f"https://api.notion.com/v1/databases/{PARAM_DB_ID}/query"
    payload = {"filter": {"property": "Method_Name", "title": {"equals": method_name}}}
    res = requests.post(url, headers=headers, json=payload)
    if res.status_code == 200 and res.json().get("results"):
        props = res.json()["results"][0]["properties"]
        def txt(n): 
            try: 
                ts = props.get(n, {}).get("rich_text", [])
                return "".join([t["text"]["content"] for t in ts]) if ts else ""
            except: return ""
        def num(n):
            try: return props.get(n, {}).get("number")
            except: return None
            
        return {
            "Instrument": txt("Instrument"), "Column_Plate": txt("Column_Plate"),
            "Condition_A": txt("Condition_A"), "Condition_B": txt("Condition_B"), "Detection": txt("Detection"),
            "SST_Criteria": txt("SST_Criteria"), "Reference_Guideline": txt("Reference_Guideline"),
            "Detail_Specificity": txt("Detail_Specificity"), "Detail_Linearity": txt("Detail_Linearity"),
            "Detail_Range": txt("Detail_Range"), "Detail_Accuracy": txt("Detail_Accuracy"),
            "Detail_Precision": txt("Detail_Precision"), "Detail_Inter_Precision": txt("Detail_Inter_Precision"),
            "Detail_LOD": txt("Detail_LOD"), "Detail_LOQ": txt("Detail_LOQ"), "Detail_Robustness": txt("Detail_Robustness"),
            "Reagent_List": txt("Reagent_List"), "Ref_Standard_Info": txt("Ref_Standard_Info"),
            "Preparation_Std": txt("Preparation_Std"), "Preparation_Sample": txt("Preparation_Sample"),
            "Calculation_Formula": txt("Calculation_Formula"), "Logic_Statement": txt("Logic_Statement"),
            "Target_Conc": num("Target_Conc"), "Unit": txt("Unit")
        }
    return {}

# ---------------------------------------------------------
# 3. 문서 생성 엔진
# ---------------------------------------------------------
def set_korean_font(doc):
    style = doc.styles['Normal']
    style.font.name = 'Malgun Gothic'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')
    style.font.size = Pt(10)

def set_font(run):
    run.font.name = 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')    

def set_table_header_style(cell):
    tcPr = cell._element.get_or_add_tcPr()
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), 'D9D9D9') 
    tcPr.append(shading_elm)
    if cell.paragraphs:
        if cell.paragraphs[0].runs: cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

# [VMP]
def generate_vmp_premium(modality, phase, df_strategy):
    doc = Document(); set_korean_font(doc)
    head = doc.add_heading('밸리데이션 종합계획서 (Validation Master Plan)', 0); head.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    table_info = doc.add_table(rows=2, cols=4); table_info.style = 'Table Grid'
    headers = ["제품명 (Product)", "단계 (Phase)", "문서 번호 (Doc No.)", "제정 일자 (Date)"]
    values = [f"{modality} Project", phase, "VMP-001", datetime.now().strftime('%Y-%m-%d')]
    for i, h in enumerate(headers): c = table_info.rows[0].cells[i]; c.text=h; set_table_header_style(c)
    for i, v in enumerate(values): c = table_info.rows[1].cells[i]; c.text=v; c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    for t, c in [("1. 목적 (Objective)", "본 계획서는 밸리데이션 전략과 범위를 규정한다."), ("2. 적용 범위 (Scope)", f"본 문서는 {modality}의 {phase} 시험법 밸리데이션에 적용된다."), ("3. 근거 가이드라인 (Reference)", "• ICH Q2(R2)\n• MFDS 가이드라인")]:
        doc.add_heading(t, level=1); doc.add_paragraph(c)
    doc.add_heading('4. 밸리데이션 수행 전략 (Validation Strategy)', level=1)
    table = doc.add_table(rows=1, cols=4); table.style = 'Table Grid'
    for i, h in enumerate(['No.', 'Method', 'Category', 'Required Items']): c = table.rows[0].cells[i]; c.text=h; set_table_header_style(c)
    for idx, row in df_strategy.iterrows(): r = table.add_row().cells; r[0].text=str(idx+1); r[1].text=str(row['Method']); r[2].text=str(row['Category']); r[3].text=", ".join(row['Required_Items'])
    doc_io = io.BytesIO(); doc.save(doc_io); doc_io.seek(0)
    return doc_io

# [Master Recipe Excel]
def generate_master_recipe_excel(method_name, target_conc, unit, stock_conc, req_vol, sample_type, powder_info=""):
    output = io.BytesIO(); workbook = xlsxwriter.Workbook(output, {'in_memory': True})
    
    # [스타일]
    title_fmt = workbook.add_format({'bold':True, 'font_size': 14, 'align':'center', 'valign':'vcenter', 'bg_color': '#44546A', 'font_color': 'white'})
    header = workbook.add_format({'bold':True, 'border':1, 'bg_color':'#D9E1F2', 'align':'center'})
    section_title = workbook.add_format({'bold':True, 'border':1, 'bg_color':'#FFC000', 'font_size':11, 'align':'left'}) 
    sub = workbook.add_format({'bold':True, 'border':1, 'bg_color':'#EDEDED', 'align':'center'})
    cell = workbook.add_format({'border':1, 'align':'center'})
    num = workbook.add_format({'border':1, 'num_format':'0.00', 'align':'center'})
    auto = workbook.add_format({'border':1, 'bg_color':'#E2EFDA', 'num_format':'0.000', 'align':'center'})
    total_fmt = workbook.add_format({'bold':True, 'border':1, 'bg_color':'#FFFF00', 'num_format':'0.00', 'align':'center'})
    
    ws = workbook.add_worksheet("Master Recipe")
    ws.set_column('A:F', 18)
    
    # [기본 정보 입력]
    ws.merge_range('A1:F1', f'Validation Material Planner: {method_name}', title_fmt)
    ws.write('A3', "Sample Type:", sub); ws.write('B3', sample_type, cell)
    if sample_type == "Powder (파우더)": ws.write('C3', "Prep Detail:", sub); ws.write_string('D3', powder_info, cell)
    ws.write('A4', "Stock Conc:", sub); ws.write('B4', stock_conc, num); ws.write('C4', unit, cell)
    ws.write('A5', "Target Conc:", sub); ws.write('B5', target_conc, num); ws.write('C5', unit, cell)
    ws.write('A6', "Vol/Vial (mL):", sub); ws.write('B6', req_vol, num)
    ws.write('D6', "TOTAL STOCK NEEDED (mL):", sub)
    
    # [희석 조제표 작성]
    ws.write(8, 0, "■ Dilution Scheme (Linearity & Accuracy)", header)
    ws.write_row(9, 0, ["Level (%)", "Target Conc", "Stock Vol (mL)", "Diluent Vol (mL)", "Total (mL)", "Check"], header)
 
    row = 10 
    start_sum_row = row + 1 # 엑셀 수식용 시작 행 (11행)

    # [공통 섹션 생성 함수]
    def add_section_grouped(main_title, levels, reps):
        nonlocal row
        ws.merge_range(row, 0, row, 5, f"■ {main_title}", header); row += 1
        data_start_row = row
        for rep in range(1, reps + 1):
            ws.merge_range(row, 0, row, 5, f"{main_title.split(' ')[0]} - {rep}회차 조제 (Set {rep})", section_title); row += 1
            ws.write_row(row, 0, ["Item ID", "Target Conc", "Stock Vol (mL)", "Diluent Vol (mL)", "Total (mL)", "Check"], sub); row += 1
            
            # 수식 범위를 위한 시작 행 저장 (현재 row는 Python index이므로 Excel row는 +1)
            start_excel_row = row + 1

            for level in levels:
                t_val = float(target_conc) * (level / 100)
                if float(stock_conc) < t_val: 
                    s_vol = "Error"
                else: 
                    s_vol = (t_val * float(req_vol)) / float(stock_conc)
                    d_vol = float(req_vol) - s_vol
                
                ws.write(row, 0, f"{main_title.split(' ')[0]}-{level}%-R{rep}", cell)
                ws.write(row, 1, t_val, num)
                
                if isinstance(s_vol, str): 
                    ws.write(row, 2, s_vol, total_fmt)
                    ws.write(row, 3, "N/A", total_fmt)
                else: 
                    ws.write(row, 2, s_vol, auto)
                    ws.write(row, 3, d_vol, auto)
                
                ws.write(row, 4, float(req_vol), num)
                ws.write(row, 5, "□", cell)
                row += 1
            
            # 수식 범위를 위한 끝 행 저장 (데이터 마지막 줄)
            end_excel_row = row

            # [수정완료] 합계 수식: C{start}:C{end}
            ws.write(row, 1, f"[{rep}회차] 소요 Stock:", sub)
            if isinstance(s_vol, str): 
                ws.write(row, 2, "Error", total_fmt)
            else: 
                ws.write_formula(row, 2, f"=SUM(C{start_excel_row}:C{end_excel_row})", total_fmt)
            row += 2

    # [섹션별 데이터 생성]        
    add_section_grouped("1. 시스템 적합성 (SST)", [100], 1)
    add_section_grouped("2. 특이성 (Specificity)", [100], 1)
    add_section_grouped("3. 직선성 (Linearity)", [80, 90, 100, 110, 120], 3)
    add_section_grouped("4. 정확성 (Accuracy)", [80, 100, 120], 3)
    
    # [정밀성 섹션]
    ws.merge_range(row, 0, row, 5, "■ 5. 정밀성 (Repeatability)", header); row += 2
    ws.merge_range(row, 0, row, 5, "반복성 시험 세트 (n=6)", section_title); row += 1
    ws.write_row(row, 0, ["Item ID", "Target Conc", "Stock Vol (mL)", "Diluent Vol (mL)", "Total (mL)", "Check"], sub); row += 1
    
    p_start_excel = row + 1 # 정밀성 데이터 시작 행 (Excel 기준)
    
    for i in range(1, 7):
        t_val = float(target_conc)
        s_vol = (t_val * float(req_vol)) / float(stock_conc)
        d_vol = float(req_vol) - s_vol
        
        ws.write(row, 0, f"Prec-100%-{i}", cell)
        ws.write(row, 1, t_val, num)
        ws.write(row, 2, s_vol, auto)
        ws.write(row, 3, d_vol, auto)
        ws.write(row, 4, float(req_vol), num)
        ws.write(row, 5, "□", cell)
        row += 1
        
    p_end_excel = row # 정밀성 데이터 끝 행
    ws.write(row, 1, "[정밀성] 소요 Stock:", sub)
    ws.write_formula(row, 2, f"=SUM(C{p_start_excel}:C{p_end_excel})", total_fmt); row += 2
    
    add_section_grouped("7. 완건성 (Robustness)", [100], 3)
    add_section_grouped("8. LOD/LOQ", [1, 0.5], 3)
    
    # 전체 합계 (단순 합산)
    ws.write_formula('E6', f"=SUM(C9:C{row})", workbook.add_format({'bold':True, 'border':1, 'bg_color':'#FF0000', 'font_color':'white', 'num_format':'0.00', 'align':'center'}))
    
    workbook.close(); output.seek(0)
    return output

# [PROTOCOL]
def generate_protocol_premium(method_name, category, params, stock_conc=None, req_vol=None, target_conc_override=None):
    doc = Document()
    
    # -----------------------------------------------------------
    # [글꼴 설정 함수] 한글: 맑은 고딕, 영어: Times New Roman
    # -----------------------------------------------------------
    def set_font(run):
        run.font.name = 'Times New Roman'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')

    # 기본 스타일 설정
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')
    style.font.size = Pt(10)
    
    # -----------------------------------------------------------
    # 1. 헤더 (Header) - 문서 번호 및 날짜 (왼쪽 정렬로 변경)
    # -----------------------------------------------------------
    section = doc.sections[0]
    header = section.header
    
    # 문서 번호 생성
    doc_no = f"VP-{method_name[:3].upper() if method_name else 'GEN'}-{datetime.now().strftime('%y%m%d')}"
    
    p_head = header.paragraphs[0]
    p_head.alignment = WD_ALIGN_PARAGRAPH.LEFT 
    r1 = p_head.add_run(f"Document No.: {doc_no}\n")
    r1.bold = True; r1.font.size = Pt(9); set_font(r1)
    r2 = p_head.add_run(f"Date: {datetime.now().strftime('%Y-%m-%d')}")
    r2.font.size = Pt(9); set_font(r2)

    # -----------------------------------------------------------
    # 2. 제목 및 개요
    # -----------------------------------------------------------
    doc.add_paragraph() 
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run('시험법 밸리데이션 상세 계획서')
    run_title.bold = True; run_title.font.size = Pt(16); set_font(run_title)
    
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = p_sub.add_run(f"(Method Validation Protocol for {method_name})")
    run_sub.font.size = Pt(12); set_font(run_sub)
    doc.add_paragraph()
    
    # 공통 헤딩 함수
    def add_custom_heading(text, level):
        p = doc.add_paragraph()
        p.style = doc.styles[f'Heading {level}']
        r = p.add_run(text)
        set_font(r)
        return p
    
    # 1. 목적
    add_custom_heading('1. 목적 (Objective)', 1)
    p = doc.add_paragraph(f"본 문서는 '{method_name}' 시험법이 의약품 품질 관리에 적합함을 검증하기 위한 구체적인 시험 절차, 시액 조제 방법 및 판정 기준을 규정한다.")
    set_font(p.runs[0])

    # 2. 기기 및 시약 (상세)
    add_custom_heading('2. 기기 및 분석 조건 (Instruments & Conditions)', 1)
    t_cond = doc.add_table(rows=5, cols=2); t_cond.style = 'Table Grid'
    cond_list = [
        ("사용 기기 (Instrument)", params.get('Instrument', 'HPLC System')),
        ("컬럼 (Column)", params.get('Column_Plate', 'C18 Column')),
        ("검출기 (Detector)", params.get('Detection', 'UV/Vis')),
        ("이동상 (Mobile Phase)", f"A: {params.get('Condition_A', 'N/A')}\nB: {params.get('Condition_B', 'N/A')}"),
        ("희석액 (Diluent)", "이동상 A와 B의 혼합액 또는 규정된 용매")
    ]
    for i, (k, v) in enumerate(cond_list):
        cell0 = t_cond.rows[i].cells[0]; cell1 = t_cond.rows[i].cells[1]
        r0 = cell0.paragraphs[0].add_run(k); r0.bold = True; set_font(r0)
        r1 = cell1.paragraphs[0].add_run(str(v)); set_font(r1)
        set_table_header_style(cell0)

    # -----------------------------------------------------------
    # 3. 항목별 상세 시험 방법 (SOP 수준 구체화)
    # -----------------------------------------------------------
    doc.add_heading('3. 상세 시험 방법 (Test Procedure)', level=1)
    
    # 변수 설정 (입력값 없으면 기본값 1.0)
    try:
        s_conc = float(stock_conc) if stock_conc else 0.0
        t_conc = float(target_conc_override) if target_conc_override else 1.0
        v_req = float(req_vol) if req_vol else 10.0
    except:
        s_conc = 0.0; t_conc = 1.0; v_req = 10.0
    unit = params.get('Unit', 'mg/mL')

    # 3.1 공통 조제 (Stock)
    add_custom_heading('3.1 시액 및 표준액 조제', 2)
    
    p_list = [
        "1) 희석액(Diluent): 이동상 A와 B를 지정된 비율로 혼합하거나 규정된 용매를 사용하여 준비한다.",
        f"2) 표준 모액(Stock Solution): 표준품을 정밀하게 달아 {s_conc} {unit} 농도가 되도록 희석액으로 녹여 조제한다.",
        f"3) 위약(Placebo): 주성분을 제외한 기제를 정밀하게 달아 {v_req} mL 부피 플라스크에 넣고 희석액으로 표선까지 채워 조제한다."
    ]
    for txt in p_list:
        p = doc.add_paragraph(txt); set_font(p.runs[0])

    # -----------------------------------------------------------
    # 3. 상세 시험 방법 (SOP 수준 - 모든 항목 계산 반영)
    # -----------------------------------------------------------
    add_custom_heading('3. 상세 시험 방법 (Test Procedure)', 1)
    
    try:
        s_conc = float(stock_conc) if stock_conc else 0.0
        t_conc = float(target_conc_override) if target_conc_override else 1.0
        v_req = float(req_vol) if req_vol else 10.0
    except: s_conc = 0.0; t_conc = 1.0; v_req = 10.0
    unit = params.get('Unit', 'mg/mL')

    # [3.1 공통 조제]
    add_custom_heading('3.1 시액 및 표준액 조제', 2)
    
    p_list = [
        "1) 희석액(Diluent): 이동상 A와 B를 지정된 비율로 혼합하거나 규정된 용매를 사용하여 준비한다.",
        f"2) 표준 모액(Stock Solution): 표준품을 정밀하게 달아 {s_conc} {unit} 농도가 되도록 희석액으로 녹여 조제한다.",
        f"3) 위약(Placebo): 주성분을 제외한 기제를 정밀하게 달아 {v_req} mL 부피 플라스크에 넣고 희석액으로 표선까지 채워 조제한다."
    ]
    for txt in p_list:
        p = doc.add_paragraph(txt); set_font(p.runs[0])

    # [3.2 특이성]
    add_custom_heading('3.2 특이성 (Specificity)', 2)
    p = doc.add_paragraph("다음 용액을 조제하여 주입한다.")
    set_font(p.runs[0])
    
    spec_list = [
        "• 공시험액, 위약: 3.1항에서 조제한 용액 사용.",
        f"• 표준액(100%): 표준 모액 {(t_conc*v_req/s_conc if s_conc>0 else 0):.3f} mL를 {v_req} mL 플라스크에 넣고 희석액으로 표선까지 채운다."
    ]
    for txt in spec_list:
        p = doc.add_paragraph(txt); set_font(p.runs[0])

    # [3.3 직선성]
    add_custom_heading('3.3 직선성 (Linearity)', 2)
    p = doc.add_paragraph(f"표준 모액({s_conc} {unit})을 사용하여 아래 표와 같이 5개 농도 레벨로 희석한다.")
    set_font(p.runs[0])
    p = doc.add_paragraph(f"※ 각 농도 레벨별로 3회씩 독립적으로 조제하여(총 15개 검액), 각각 1회 분석한다.")
    set_font(p.runs[0])
    
    t_lin = doc.add_table(rows=1, cols=5); t_lin.style = 'Table Grid'
    headers = ["Level", "목표 농도", "모액 취함 (mL)", "최종 부피 (mL)", "희석액 (mL)"]
    for i, h in enumerate(headers): 
        c = t_lin.rows[0].cells[i]
        r = c.paragraphs[0].add_run(h); r.bold = True; set_font(r)
        set_table_header_style(c)

    for level in [80, 90, 100, 110, 120]:
        row = t_lin.add_row().cells
        tgt = t_conc * (level/100)
        vs = (tgt * v_req) / s_conc if s_conc > 0 else 0
        vd = v_req - vs
        
        row[0].text = f"{level}%"
        row[1].text = f"{tgt:.4f} {unit}"
        row[2].text = f"{vs:.3f}"
        row[3].text = f"{v_req:.1f}"
        row[4].text = f"{vd:.3f}"
        for cell in row: set_font(cell.paragraphs[0].runs[0])

    # [3.4 정확성]
    add_custom_heading('3.4 정확성 (Accuracy)', 2)
    p = doc.add_paragraph("기준 농도의 80%, 100%, 120% 수준으로 각 3회씩 독립적으로 조제하여 분석한다 (총 9개 검액).")
    set_font(p.runs[0])
    
    acc_list = [
        f"• 80% Level (3회): 위 직선성 표의 80% 조건({(t_conc*0.8*v_req/s_conc):.3f} mL 모액 → {v_req} mL)으로 3개 조제.",
        f"• 100% Level (3회): 위 직선성 표의 100% 조건({(t_conc*1.0*v_req/s_conc):.3f} mL 모액 → {v_req} mL)으로 3개 조제.",
        f"• 120% Level (3회): 위 직선성 표의 120% 조건({(t_conc*1.2*v_req/s_conc):.3f} mL 모액 → {v_req} mL)으로 3개 조제."
    ]
    for txt in acc_list:
        p = doc.add_paragraph(txt); set_font(p.runs[0])

    # [3.5 정밀성]
    add_custom_heading('3.5 정밀성 (Precision)', 2)
    p = doc.add_paragraph(f"기준 농도(100%)인 {t_conc} {unit} 검액을 6개 독립적으로 조제한다.")
    set_font(p.runs[0])
    p = doc.add_paragraph(f"• 조제법: 표준 모액 {(t_conc*v_req/s_conc):.3f} mL를 취하여 {v_req} mL 부피 플라스크에 넣고 희석한다. (x 6회 반복)")
    set_font(p.runs[0])

    # [3.6 LOD/LOQ] - 중간 희석액 도입
    add_custom_heading('3.6 검출 및 정량한계 (LOD/LOQ)', 2)
    p = doc.add_paragraph("저농도에서의 정확한 조제를 위해 '중간 희석액'을 거쳐 단계적으로 희석한다.")
    set_font(p.runs[0])
    
    # 중간 희석액 계산 (타겟의 10% 수준)
    inter_conc = t_conc * 0.1
    inter_vol_req = 100.0 # 중간 희석액은 넉넉하게 100mL 제조 가정
    stock_for_inter = (inter_conc * inter_vol_req) / s_conc if s_conc > 0 else 0
    
    p = doc.add_paragraph(f"1) 중간 희석액 조제 ({inter_conc:.4f} {unit}): 표준 모액 {stock_for_inter:.3f} mL를 취하여 {inter_vol_req} mL 부피 플라스크에 넣고 희석한다.")
    set_font(p.runs[0])
    
    t_lod = doc.add_table(rows=1, cols=5); t_lod.style = 'Table Grid'
    lh = ["구분", "추정 Level", "농도", "중간액 취함 (mL)", "최종 부피 (mL)"]
    for i, h in enumerate(lh): 
        c = t_lod.rows[0].cells[i]
        r = c.paragraphs[0].add_run(h); r.bold = True; set_font(r)
        set_table_header_style(c)
    
    # LOQ (1%), LOD (0.3% 가정)
    for lvl, name in [(1.0, "LOQ (예상)"), (0.33, "LOD (예상)")]:
        lr = t_lod.add_row().cells
        ltgt = t_conc * (lvl/100)
        # 중간액에서 희석: V = (Target * Total) / Inter_Conc
        lvs = (ltgt * v_req) / inter_conc if inter_conc > 0 else 0
        
        lr[0].text = name; lr[1].text = f"{lvl}%"; lr[2].text = f"{ltgt:.5f}"
        lr[3].text = f"{lvs:.3f}"; lr[4].text = f"{v_req:.1f}"
        for c in lr: set_font(c.paragraphs[0].runs[0])

    # -----------------------------------------------------------
    # 4. 밸리데이션 항목 및 판정 기준 (서술식 & 분리)
    # -----------------------------------------------------------
    doc.add_heading('4. 밸리데이션 항목 및 판정 기준 (Evaluation & Criteria)', level=1)
    
    # 4.1 특이성
    doc.add_heading('4.1 특이성 (Specificity)', level=2)
    doc.add_paragraph("1) 평가 방법 (Evaluation Method)")
    doc.add_paragraph("   공시험액(Blank), 위약(Placebo), 표준액을 각각 분석하여 크로마토그램을 비교한다. 주성분 피크의 머무름 시간(RT)에 간섭하는 피크가 있는지 확인한다.")
    doc.add_paragraph("2) 판정 기준 (Acceptance Criteria)")
    crit_spec = params.get('Detail_Specificity', "간섭 피크 면적 ≤ 표준액 평균 면적의 0.5%")
    doc.add_paragraph(f"   - 공시험액 및 위약에서 주성분 피크와 겹치는 간섭 피크가 없거나, 검출되더라도 그 면적이 {crit_spec} 이어야 한다.")

    # 4.2 직선성
    doc.add_heading('4.2 직선성 (Linearity)', level=2)
    doc.add_paragraph("1) 평가 방법 (Evaluation Method)")
    doc.add_paragraph(f"   {t_conc} {unit} 농도를 기준으로 80 ~ 120% 범위 내 5개 농도의 표준액을 분석한다. 농도(X축)와 피크 면적(Y축)에 대한 회귀분석을 수행하여 상관계수(R) 및 결정계수(R²)를 구한다.")
    doc.add_paragraph("2) 판정 기준 (Acceptance Criteria)")
    crit_lin = params.get('Detail_Linearity', "결정계수(R²) ≥ 0.990")
    doc.add_paragraph(f"   - {crit_lin}")
    doc.add_paragraph("   - Y절편과 기울기가 타당한 수준이어야 한다.")

    # 4.3 정확성
    doc.add_heading('4.3 정확성 (Accuracy)', level=2)
    doc.add_paragraph("1) 평가 방법 (Evaluation Method)")
    doc.add_paragraph("   기준 농도의 80%, 100%, 120% 수준에서 각각 3회씩 조제하여 분석한다. 각 검액의 실측 농도를 이론 농도로 나누어 회수율(Recovery, %)을 계산한다.")
    doc.add_paragraph("2) 판정 기준 (Acceptance Criteria)")
    crit_acc = params.get('Detail_Accuracy', "회수율 80.0 ~ 120.0%")
    doc.add_paragraph(f"   - 각 농도별 평균 회수율 및 전체 평균 회수율이 {crit_acc} 이내여야 한다.")
    doc.add_paragraph("   - 각 농도별 회수율의 상대표준편차(RSD)가 적절해야 한다.")

    # 4.4 정밀성
    doc.add_heading('4.4 정밀성 (Precision)', level=2)
    doc.add_paragraph("1) 평가 방법 (Evaluation Method)")
    doc.add_paragraph("   기준 농도(100%)에 해당하는 검액을 6개 독립적으로 조제하여 분석한다. 6회 결과에 대한 피크 면적의 상대표준편차(RSD)를 계산한다.")
    doc.add_paragraph("2) 판정 기준 (Acceptance Criteria)")
    crit_prec = params.get('Detail_Precision', "RSD ≤ 2.0%")
    doc.add_paragraph(f"   - 피크 면적의 {crit_prec}")

    # 4.5 정량한계
    doc.add_heading('4.5 검출 및 정량한계 (LOD & LOQ)', level=2)
    doc.add_paragraph("1) 평가 방법 (Evaluation Method)")
    doc.add_paragraph("   신호 대 잡음비(Signal-to-Noise Ratio, S/N) 방식을 이용한다. 예상되는 저농도 용액을 분석하여 S/N 비를 측정한다.")
    doc.add_paragraph("2) 판정 기준 (Acceptance Criteria)")
    crit_loq = params.get('Detail_LOQ', "LOD S/N ≥ 3, LOQ S/N ≥ 10")
    doc.add_paragraph(f"   - {crit_loq}")

    # -----------------------------------------------------------
    # 5. 서명
    # -----------------------------------------------------------
    doc.add_paragraph("\n\n")
    t_sign = doc.add_table(rows=2, cols=3); t_sign.style = 'Table Grid'
    roles = ["작성자 (Prepared By)", "검토자 (Reviewed By)", "승인자 (Approved By)"]
    for i, r in enumerate(roles): 
        c = t_sign.rows[0].cells[i]; c.text = r; set_table_header_style(c)
        t_sign.rows[1].cells[i].text = "\n\n서명: _______________\n날짜: _______________\n"

    doc_io = io.BytesIO(); doc.save(doc_io); doc_io.seek(0)
    return doc_io

# [Excel 생성 함수 - Smart Logbook (ACTUAL WEIGHT & CORRECTION LOGIC)]
def generate_smart_excel(method_name, category, params, simulate=False):
    output = io.BytesIO()
    workbook = xlsxwriter.Workbook(output, {'in_memory': True})
    
    # Styles
    header = workbook.add_format({'bold':True, 'border':1, 'bg_color':'#4472C4', 'font_color':'white', 'align':'center', 'valign':'vcenter'})
    sub = workbook.add_format({'bold':True, 'border':1, 'bg_color':'#D9E1F2', 'align':'center'})
    cell = workbook.add_format({'border':1, 'align':'center'})
    num = workbook.add_format({'border':1, 'num_format':'0.00', 'align':'center'})
    num3 = workbook.add_format({'border':1, 'num_format':'0.000', 'align':'center'}) 
    auto = workbook.add_format({'border':1, 'bg_color':'#E2EFDA', 'align':'center'})
    pass_fmt = workbook.add_format({'bold':True, 'bg_color':'#C6EFCE', 'font_color':'#006100', 'align':'center'})
    fail_fmt = workbook.add_format({'bold':True, 'bg_color':'#FFC7CE', 'font_color':'#9C0006', 'align':'center'})
    crit_fmt = workbook.add_format({'bold':True, 'font_color':'red', 'align':'left'})
    
    # 1. Info Sheet (Enhanced with Actual Weighing & Purity)
    ws1 = workbook.add_worksheet("1. Info"); ws1.set_column('A:A', 25); ws1.set_column('B:E', 15); ws1.merge_range('A1:E1', f'GMP Logbook: {method_name}', header)
    info = [("Date", datetime.now().strftime("%Y-%m-%d")), ("Instrument", params.get('Instrument')), ("Column", params.get('Column_Plate')), ("Analyst", "")]
    r = 3; 
    for k, v in info: ws1.write(r, 0, k, sub); ws1.merge_range(r, 1, r, 4, v if v else "", cell); r+=1
    ws1.write(r+1, 0, "Round Rule:", sub); ws1.merge_range(r+1, 1, r+1, 4, "모든 계산값은 소수점 2째자리에서 절사(ROUNDDOWN)함.", cell)
    
    # Actual Stock Prep Section
    r += 3
    ws1.merge_range(r, 0, r, 4, "■ Standard Stock Solution Preparation (보정값 적용)", sub_rep); r+=1
    ws1.write(r, 0, "Purity (Potency, %):", sub); ws1.write(r, 1, "", calc); ws1.write(r, 2, "%", cell)
    ws1.write(r+1, 0, "Water Content (%):", sub); ws1.write(r+1, 1, 0, calc); ws1.write(r+1, 2, "% (If applicable)", cell)
    ws1.write(r+2, 0, "Actual Weight (mg):", sub); ws1.write(r+2, 1, "", calc); ws1.write(r+2, 2, "mg", cell)
    ws1.write(r+3, 0, "Final Volume (mL):", sub); ws1.write(r+3, 1, "", calc); ws1.write(r+3, 2, "mL", cell)
    ws1.write(r+4, 0, "Actual Stock Conc (mg/mL):", sub)
    # Actual Conc = (Weight * (Purity/100) * ((100-Water)/100)) / Vol
    # Assuming B11=Purity, B12=Water, B13=Weight, B14=Vol
    # Formula Row Index: r is variable. Purity at r, Weight at r+2.
    purity_cell = f"B{r+1}"; water_cell = f"B{r+2}"; weight_cell = f"B{r+3}"; vol_cell = f"B{r+4}"
    ws1.write_formula(r+4, 1, f"=ROUNDDOWN(({weight_cell}*({purity_cell}/100)*((100-{water_cell})/100))/{vol_cell}, 4)", total_fmt)
    actual_stock_ref = f"'1. Info'!B{r+5}" # Reference for other sheets

    # 2. SST Sheet
    ws_sst = workbook.add_worksheet("2. SST"); ws_sst.set_column('A:F', 15)
    ws_sst.merge_range('A1:F1', 'System Suitability Test (n=6)', header)
    ws_sst.write_row('A2', ["Inj No.", "RT (min)", "Area", "Height", "Tailing (1st)", "Plate Count"], sub)
    for i in range(1, 7): ws_sst.write(i+1, 0, i, cell); ws_sst.write_row(i+1, 1, ["", "", "", "", ""], calc)
    ws_sst.write('A9', "Mean", sub); ws_sst.write_formula('B9', "=ROUNDDOWN(AVERAGE(B3:B8), 2)", auto); ws_sst.write_formula('C9', "=ROUNDDOWN(AVERAGE(C3:C8), 2)", auto)
    ws_sst.write('A10', "RSD(%)", sub); ws_sst.write_formula('B10', "=ROUNDDOWN(STDEV(B3:B8)/B9*100, 2)", auto); ws_sst.write_formula('C10', "=ROUNDDOWN(STDEV(C3:C8)/C9*100, 2)", auto)
    ws_sst.write('A12', "Criteria (RSD):", sub); ws_sst.write('B12', "≤ 2.0%", cell)
    ws_sst.write('C12', "Criteria (Tail):", sub); ws_sst.write('D12', "≤ 2.0 (Inj #1)", cell) 
    ws_sst.write('E12', "Result:", sub)
    ws_sst.write_formula('F12', '=IF(AND(B10<=2.0, C10<=2.0, E3<=2.0), "Pass", "Fail")', pass_fmt)
    ws_sst.conditional_format('F12', {'type': 'cell', 'criteria': '==', 'value': '"Fail"', 'format': fail_fmt})

    # [Criteria Added]
    ws_sst.write('A14', "※ Acceptance Criteria:", crit_fmt)
    ws_sst.write('A15', "1) RSD of RT & Area ≤ 2.0%")
    ws_sst.write('A16', "2) Tailing Factor (1st Inj) ≤ 2.0")

    # 3. Specificity Sheet
    ws_spec = workbook.add_worksheet("3. Specificity"); ws_spec.set_column('A:E', 20)
    ws_spec.merge_range('A1:E1', 'Specificity Test (Identification & Interference)', header)
    
    # [Reference Data from SST] - SST 결과값 자동 참조
    ws_spec.write('A3', "Ref. Std RT (min):", sub); ws_spec.write_formula('B3', "='2. SST'!B9", num) # SST Mean RT
    ws_spec.write('C3', "Ref. Std Area:", sub); ws_spec.write_formula('D3', "='2. SST'!C9", num) # SST Mean Area
    
    # -----------------------------------------------------------
    # Part 1. Identification (RT Match) - 주성분 확인
    # -----------------------------------------------------------
    ws_spec.merge_range('A5:E5', "1. Identification (RT Match)", sub_rep)
    ws_spec.write_row('A6', ["Sample", "RT (min)", "Diff with Std (%)", "Criteria (≤2.0%)", "Result"], sub)
    
    # 검체(Sample) 1개 예시
    ws_spec.write('A7', "Sample", cell)
    ws_spec.write('B7', "", calc) # 사용자 입력 (검체 RT)
    
    # RT 차이(%) = abs(검체RT - 표준RT) / 표준RT * 100
    ws_spec.write_formula('C7', f"=IF(B7=\"\",\"\",ROUNDDOWN(ABS(B7-$B$3)/$B$3*100, 2))", auto)
    ws_spec.write('D7', "≤ 2.0%", cell)
    ws_spec.write_formula('E7', f'=IF(C7=\"\",\"\",IF(C7<=2.0, "Pass", "Fail"))', pass_fmt)
    ws_spec.conditional_format('E7', {'type': 'cell', 'criteria': '==', 'value': '"Fail"', 'format': fail_fmt})

    # -----------------------------------------------------------
    # Part 2. Interference (Area Check) - 간섭 확인
    # -----------------------------------------------------------
    ws_spec.merge_range('A9:E9', "2. Interference (Blank/Placebo Check)", sub_rep)
    ws_spec.write_row('A10', ["Sample", "Detected RT", "Area", "Interference (%)", "Result (≤0.5%)"], sub)
    
    for i, s in enumerate(["Blank", "Placebo"]):
        row = i + 11
        ws_spec.write(row, 0, s, cell)
        ws_spec.write(row, 1, "", calc) # RT 입력 (간섭 피크가 떴을 때)
        ws_spec.write(row, 2, "", calc) # Area 입력
        
        # 간섭율(%) = (간섭피크 면적 / 표준액 평균 면적) * 100
        # 분모(D3)가 0이거나 비어있을 때 에러 방지
        ws_spec.write_formula(row, 3, f"=IF(OR($D$3=\"\",$D$3=0), \"\", IF(C{row+1}=\"\", 0, ROUNDDOWN(C{row+1}/$D$3*100, 2)))", auto)
        
        # 판정: 0.5% 이하 Pass
        ws_spec.write_formula(row, 4, f'=IF(D{row+1}<=0.5, "Pass", "Fail")', pass_fmt)
        ws_spec.conditional_format(f'E{row+1}', {'type': 'cell', 'criteria': '==', 'value': '"Fail"', 'format': fail_fmt})

    # [Criteria Added]
    ws_spec.write(14, 0, "※ Acceptance Criteria:", crit_fmt)
    ws_spec.write(15, 0, "1) Interference Peak Area ≤ 0.5% of Standard Area")

    # 4. Linearity Sheet (Uses Actual Stock Conc)
    target_conc = params.get('Target_Conc')
    if target_conc:
        ws2 = workbook.add_worksheet("4. Linearity"); ws2.set_column('A:I', 13)
        unit = params.get('Unit', 'ppm'); ws2.merge_range('A1:I1', f'Linearity Test (Target: {target_conc} {unit})', header)
        row = 3; rep_rows = {1: [], 2: [], 3: []}
        
        for rep in range(1, 4):
            ws2.merge_range(row, 0, row, 8, f"■ Repetition {rep}", sub_rep); row += 1
            ws2.write_row(row, 0, ["Level", "Conc (X)", "Area (Y)", "Back Calc", "Accuracy (%)", "Check"], sub); row += 1
            data_start = row
            for level in [80, 90, 100, 110, 120]:
                # Conc (X) now links to Info Sheet Actual Stock * (Level/100) or similar dilution logic
                # Assuming simple dilution from stock: Actual Stock * (Level % of Target / Stock?) -> This depends on recipe.
                # Simplified: Actual Stock * (Target * Level% / Stock_Target_Ratio)
                # Let's assume standard dilution: X = Actual_Stock * (Level/100) if Stock was made to be 100%. 
                # But stock is usually hi-conc. Let's assume the user prepared levels to match 80%~120% of TARGET.
                # So Conc X = Target_Conc_Theoretical * (Actual_Stock / Theoretical_Stock) * Level%
                # Ideally, simple reference: =Actual_Stock_Cell * Dilution_Factor
                # For this template, we will allow user to input Actual Conc X or calc from Info.
                # Best approach: X = Actual Stock * (Level_Target / Stock_Target)
                ws2.write(row, 0, f"{level}%", cell)
                # Here we simply assume they diluted to nominal targets relative to the actual stock
                # Formula: =Info!ActualStock * (Level/100) * (Target/Stock_User_Input) -> Complex.
                # Use simplified: =ROUNDDOWN(ActualStock * (Level/100), 3) assuming Stock is ~100% target or normalized.
                # Let's link to the calculated actual stock from Info sheet as base
                ws2.write_formula(row, 1, f"=ROUNDDOWN({actual_stock_ref} * ({level}/100), 3)", num) # Dynamic Actual Conc
                ws2.write(row, 2, "", calc)
                rep_rows[rep].append(row + 1)
                ind_slope = f"C{data_start+7}"; ind_int = f"C{data_start+8}"
                ws2.write_formula(row, 3, f"=IF(C{row+1}<>\"\", ROUNDDOWN((C{row+1}-{ind_int})/{ind_slope}, 3), \"\")", auto)
                ws2.write_formula(row, 4, f"=IF(C{row+1}<>\"\", ROUNDDOWN(D{row+1}/B{row+1}*100, 1), \"\")", auto)
                ws2.write(row, 5, "OK", cell); row += 1
            ws2.write(row, 1, "Slope:", sub); ws2.write_formula(row, 2, f"=SLOPE(C{data_start+1}:C{row}, B{data_start+1}:B{row})", auto)
            ws2.write(row+1, 1, "Intercept:", sub); ws2.write_formula(row+1, 2, f"=INTERCEPT(C{data_start+1}:C{row}, B{data_start+1}:B{row})", auto)
            ws2.write(row+2, 1, "R²:", sub); ws2.write_formula(row+2, 2, f"=RSQ(C{data_start+1}:C{row}, B{data_start+1}:B{row})", auto)
            chart = workbook.add_chart({'type': 'scatter', 'subtype': 'straight_with_markers'})
            chart.add_series({'name': f'Rep {rep}', 'categories': f"='4. Linearity'!$B${data_start+1}:$B${row}", 'values': f"='4. Linearity'!$C${data_start+1}:$C${row}", 'trendline': {'type': 'linear', 'display_equation': True, 'display_r_squared': True}})
            chart.set_size({'width': 350, 'height': 220}); ws2.insert_chart(f'G{data_start}', chart)
            row += 6

        ws2.merge_range(row, 0, row, 8, "■ Summary (Mean of 3 Reps) & Final Check", sub_rep); row += 1
        ws2.write_row(row, 0, ["Level", "Conc (X)", "Mean Area", "STDEV", "% RSD", "Criteria (RSD≤5%)"], sub); row += 1
        summary_start = row
        for i, level in enumerate([80, 90, 100, 110, 120]):
            r1 = rep_rows[1][i]; r2 = rep_rows[2][i]; r3 = rep_rows[3][i]
            ws2.write(row, 0, f"{level}%", cell); ws2.write_formula(row, 1, f"=B{r1}", num)
            ws2.write_formula(row, 2, f"=ROUNDDOWN(AVERAGE(C{r1},C{r2},C{r3}), 2)", auto)
            ws2.write_formula(row, 3, f"=ROUNDDOWN(STDEV(C{r1},C{r2},C{r3}), 2)", auto)
            ws2.write_formula(row, 4, f"=ROUNDDOWN(IF(C{row+1}=0, 0, D{row+1}/C{row+1}*100), 2)", auto)
            ws2.write_formula(row, 5, f'=IF(E{row+1}<=5.0, "Pass", "Fail")', pass_fmt)
            row += 1
        row += 1
        slope_cell = f"'4. Linearity'!C{row+1}"; int_cell = f"'4. Linearity'!C{row+2}"
        ws2.write(row, 1, "Slope:", sub); ws2.write_formula(row, 2, f"=ROUNDDOWN(SLOPE(C{summary_start+1}:C{summary_start+5}, B{summary_start+1}:B{summary_start+5}), 4)", auto)
        ws2.write(row+1, 1, "Intercept:", sub); ws2.write_formula(row+1, 2, f"=ROUNDDOWN(INTERCEPT(C{summary_start+1}:C{summary_start+5}, B{summary_start+1}:B{summary_start+5}), 4)", auto)
        ws2.write(row+2, 1, "R²:", sub); ws2.write_formula(row+2, 2, f"=ROUNDDOWN(RSQ(C{summary_start+1}:C{summary_start+5}, B{summary_start+1}:B{summary_start+5}), 4)", auto)
        ws2.write(row+2, 3, "Criteria (≥0.990):", sub); ws2.write_formula(row+2, 4, f'=IF(C{row+3}>=0.990, "Pass", "Fail")', pass_fmt)

    # [Criteria Added]
    ws2.write(row+4, 0, "※ Acceptance Criteria:", crit_fmt)
    ws2.write(row+5, 0, "1) Coefficient of determination (R²) ≥ 0.990")
    ws2.write(row+6, 0, "2) %RSD of peak areas at each level ≤ 5.0%")

    # 5. Accuracy Sheet
    ws_acc = workbook.add_worksheet("5. Accuracy"); ws_acc.set_column('A:G', 15)
    ws_acc.merge_range('A1:G1', 'Accuracy (Recovery)', header)
    
    # Reference Linearity Slope/Int
    ws_acc.write('E3', "Slope:", sub); ws_acc.write_formula('F3', f"='4. Linearity'!C{row+1}", auto)
    ws_acc.write('E4', "Int:", sub); ws_acc.write_formula('F4', f"='4. Linearity'!C{row+2}", auto)
    ws_acc.write('G3', "(From Linearity)", cell)
    
    acc_row = 6
    for level in [80, 100, 120]:
        ws_acc.merge_range(acc_row, 0, acc_row, 6, f"■ Level {level}% (3 Reps)", sub_rep); acc_row += 1
        ws_acc.write_row(acc_row, 0, ["Rep", "Theo Conc", "Area", "Calc Conc", "Recovery (%)", "Criteria", "Result"], sub); acc_row += 1
        start_r = acc_row
        for rep in range(1, 4):
            ws_acc.write(acc_row, 0, rep, cell)
            # Theo Conc Formula
            ws_acc.write_formula(acc_row, 1, f"=ROUNDDOWN({actual_stock_ref} * ({level}/100), 3)", num3)
            ws_acc.write(acc_row, 2, "", calc)
            ws_acc.write_formula(acc_row, 3, f'=IF(C{acc_row+1}="","",ROUNDDOWN((C{acc_row+1}-$F$4)/$F$3, 3))', auto)
            ws_acc.write_formula(acc_row, 4, f'=IF(D{acc_row+1}="","",ROUNDDOWN(D{acc_row+1}/B{acc_row+1}*100, 1))', auto)
            ws_acc.write(acc_row, 5, "80~120%", cell)
            ws_acc.write_formula(acc_row, 6, f'=IF(E{acc_row+1}="","",IF(AND(E{acc_row+1}>=80, E{acc_row+1}<=120), "Pass", "Fail"))', pass_fmt)
            ws_acc.conditional_format(f'G{acc_row+1}', {'type': 'cell', 'criteria': '==', 'value': '"Fail"', 'format': fail_fmt}); acc_row += 1
        ws_acc.write(acc_row, 3, "Mean Rec(%):", sub); ws_acc.write_formula(acc_row, 4, f"=ROUNDDOWN(AVERAGE(E{start_r+1}:E{acc_row}), 1)", total_fmt); acc_row += 2
    
    # [Criteria Added]
    ws_acc.write(acc_row, 0, "※ Acceptance Criteria:", crit_fmt)
    ws_acc.write(acc_row+1, 0, "1) Individual & Mean Recovery: 80.0 ~ 120.0%")

    # 6. Precision, 7. Robustness, 8. LOD/LOQ (Same as before)
    ws3 = workbook.add_worksheet("6. Precision"); ws3.set_column('A:E', 15); ws3.merge_range('A1:E1', 'Precision', header)
    ws3.merge_range('A3:E3', "■ Day 1 (Repeatability)", sub); ws3.write_row('A4', ["Inj", "Sample", "Result", "Mean", "RSD"], sub)
    for i in range(6): ws3.write_row(4+i, 0, [i+1, "Sample", ""], calc)
    ws3.write_formula('D5', "=ROUNDDOWN(AVERAGE(C5:C10), 2)", num); ws3.write_formula('E5', "=ROUNDDOWN(STDEV(C5:C10)/D5*100, 2)", num)
    ws3.write('E11', "Check (RSD≤2.0):", sub); ws3.write_formula('E12', '=IF(E5<=2.0, "Pass", "Fail")', pass_fmt)
    ws3.merge_range('A14:E14', "■ Day 2 (Intermediate Precision)", sub); ws3.write_row('A15', ["Inj", "Sample", "Result", "Mean", "RSD"], sub)
    for i in range(6): ws3.write_row(15+i, 0, [i+1, "Sample", ""], calc)
    ws3.write_formula('D16', "=ROUNDDOWN(AVERAGE(C16:C21), 2)", num); ws3.write_formula('E16', "=ROUNDDOWN(STDEV(C16:C21)/D16*100, 2)", num)
    ws3.write('A23', "Diff (%)", sub); ws3.write_formula('B23', "=ROUNDDOWN(ABS(D5-D16)/AVERAGE(D5,D16)*100, 2)", num)

    if params.get('Detail_Robustness'):
        ws4 = workbook.add_worksheet("7. Robustness"); ws4.set_column('A:F', 18); ws4.merge_range('A1:F1', 'Robustness Conditions', header)
        ws4.write_row('A3', ["Condition", "Set", "Actual", "SST Result", "Pass/Fail", "Note"], sub)
        for r, c in enumerate(["Standard", "Flow -0.1", "Flow +0.1", "Temp -2", "Temp +2"]): 
            ws4.write(4+r, 0, c, cell); ws4.write_row(4+r, 1, [""]*5, calc)

    ws_ll = workbook.add_worksheet("8. LOD_LOQ"); ws_ll.set_column('A:E', 15); ws_ll.merge_range('A1:E1', 'LOD / LOQ', header)
    ws_ll.write_row('A2', ["Item", "Signal", "Noise", "S/N Ratio", "Result"], sub)
    ws_ll.write('A3', "LOD Sample", cell); ws_ll.write('B3', "", calc); ws_ll.write('C3', "", calc); ws_ll.write_formula('D3', "=ROUNDDOWN(B3/C3, 1)", auto)
    ws_ll.write_formula('E3', '=IF(D3>=3, "Pass", "Fail")', pass_fmt)
    ws_ll.write('A4', "LOQ Sample", cell); ws_ll.write('B4', "", calc); ws_ll.write('C4', "", calc); ws_ll.write_formula('D4', "=ROUNDDOWN(B4/C4, 1)", auto)
    ws_ll.write_formula('E4', '=IF(D4>=10, "Pass", "Fail")', pass_fmt)

    workbook.close(); output.seek(0)
    return output

# [누락된 함수 추가] 엑셀 데이터 추출 함수
def extract_logbook_data(uploaded_file):
    results = {}
    try:
        # 1. SST 결과
        df_sst = pd.read_excel(uploaded_file, sheet_name='2. SST', header=None)
        res_row = df_sst[df_sst.eq("Result:").any(axis=1)].index
        if not res_row.empty:
            results['sst'] = df_sst.iloc[res_row[0], 5]  # F열 값
        
        # 2. 직선성 결과 (R²)
        df_lin = pd.read_excel(uploaded_file, sheet_name='4. Linearity', header=None)
        r2_row = df_lin[df_lin.eq("Final R²:").any(axis=1)].index
        if not r2_row.empty:
            results['r2'] = df_lin.iloc[r2_row[0], 2]  # C열 값

        # 3. 정확성/정밀성 등은 필요 시 추가 파싱
        # (단순화를 위해 성공 시 기본값 반환)
        if 'sst' in results:
            return results
        else:
            return {'sst': 'N/A', 'r2': 'N/A'}
            
    except Exception as e:
        return {'error': str(e)}
    return {}

# ---------------------------------------------------------
# 4. 메인 UI
# ---------------------------------------------------------
st.set_page_config(page_title="AtheraCLOUD Full GMP", layout="wide")
st.title("🧪 AtheraCLOUD: Full CMC Validation Suite")
st.markdown("##### Strategy · Protocol · Multi-Sheet Logbook · Report")

col1, col2 = st.columns([1, 3])
with col1:
    st.header("📂 Project")
    sel_modality = st.selectbox("Modality", ["mAb", "Cell Therapy"])
    sel_phase = st.selectbox("Phase", ["Phase 1", "Phase 3"])

with col2:
    try: criteria_map = get_criteria_map(); df_full = get_strategy_list(criteria_map)
    except: df_full = pd.DataFrame()

    if sel_modality == "mAb" and not df_full.empty:
        my_plan = df_full[(df_full["Modality"] == sel_modality) & (df_full["Phase"] == sel_phase)]
        if not my_plan.empty:
            t1, t2, t3 = st.tabs(["📑 Step 1: Strategy & Protocol", "📗 Step 2: Excel Logbook", "📊 Step 3: Result Report"])
            
            with t1:
                st.markdown("### 1️⃣ 전략 (VMP) 및 상세 계획서 (Protocol)")
                st.dataframe(my_plan[["Method", "Category"]])
                c1, c2 = st.columns(2)
                with c1: st.download_button("📥 VMP(종합계획서) 다운로드", generate_vmp_premium(sel_modality, sel_phase, my_plan), "VMP_Master.docx")
                with c2:
                    st.divider()
                    st.markdown("#### 🧪 시약 제조 및 계획서 생성기")
                    sel_p = st.selectbox("Protocol:", my_plan["Method"].unique())
                    if sel_p:
                        st.info("👇 시료 상태와 농도를 입력하세요. (Target 농도가 100% 기준이 됩니다)")
                        sample_type = st.radio("시료 타입 (Sample Type):", ["Liquid (액체)", "Powder (파우더)"], horizontal=True)
                        cc1, cc2 = st.columns(2)
                        stock_input_val = 0.0; powder_desc = ""
                        if sample_type == "Liquid (액체)":
                            with cc1: stock_input_val = st.number_input("내 Stock 농도 (mg/mL 등):", min_value=0.0, step=0.1, format="%.2f")
                        else: 
                            with cc1: weight_input = st.number_input("칭량값 (Weight, mg):", min_value=0.0, step=0.1)
                            with cc2: dil_vol_input = st.number_input("희석 부피 (Vol, mL):", min_value=0.1, value=10.0, step=1.0)
                            if dil_vol_input > 0:
                                stock_input_val = weight_input / dil_vol_input
                                st.caption(f"🧪 계산된 Stock 농도: **{stock_input_val:.2f} mg/mL**")
                                powder_desc = f"Weigh {weight_input}mg / {dil_vol_input}mL"
                        params_p = get_method_params(sel_p); db_target = params_p.get('Target_Conc', 0.0)
                        with cc1: target_input_val = st.number_input("기준 농도 (Target 100%, mg/mL):", min_value=0.001, value=float(db_target) if db_target else 1.0, format="%.3f")
                        with cc2: vol_input = st.number_input("개별 바이알 조제 목표량 (Target Vol, mL):", min_value=1.0, value=5.0, step=1.0)
                        unit_val = params_p.get('Unit', '')
                        if stock_input_val > 0 and target_input_val > 0:
                            if stock_input_val < target_input_val * 1.2: st.error("⚠️ Stock 농도가 Target 농도(120% 범위)보다 낮습니다! 더 진한 Stock을 준비하세요.")
                            else:
                                calc_excel = generate_master_recipe_excel(sel_p, target_input_val, unit_val, stock_input_val, vol_input, sample_type, powder_desc)
                                st.download_button("🧮 시약 제조 계산기 (Master Recipe) 다운로드", calc_excel, f"Master_Recipe_{sel_p}.xlsx")
                                doc_proto = generate_protocol_premium(sel_p, "Cat", params_p, stock_input_val, vol_input, target_input_val)
                                st.download_button("📄 상세 계획서 (Protocol) 다운로드", doc_proto, f"Protocol_{sel_p}.docx", type="primary")

            with t2:
                st.markdown("### 📗 스마트 엑셀 일지 (Final Fixed)")
                st.info("✅ SST(Tailing Check), 특이성(Std 기준), 직선성(회차별 그래프), 정확성(자동 참조) 기능 탑재")
                sel_l = st.selectbox("Logbook:", my_plan["Method"].unique(), key="l")
                if st.button("Download Excel Logbook"):
                    data = generate_smart_excel(sel_l, "Cat", get_method_params(sel_l))
                    st.download_button("📊 Excel Logbook 다운로드", data, f"Logbook_{sel_l}.xlsx")

            with t3:
                st.markdown("### 📊 최종 결과 보고서")
                st.info("작성된 엑셀 파일을 업로드하면 결과가 자동 반영됩니다.")
                uploaded_log = st.file_uploader("📂 Upload Filled Logbook", type=["xlsx"])
                sel_r = st.selectbox("Report for:", my_plan["Method"].unique(), key="r")
                
                if uploaded_log:
                    data = extract_logbook_data(uploaded_log)
                    st.success("데이터 추출 완료!")
                    st.json(data)
                    if st.button("Generate Final Report"):
                        doc = generate_summary_report_gmp(sel_r, "Cat", get_method_params(sel_r), {'lot': 'Test'}, data)
                        st.download_button("📥 Download Report", doc, "Final_Report.docx")