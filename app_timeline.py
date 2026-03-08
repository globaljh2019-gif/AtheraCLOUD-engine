import xlsxwriter
import streamlit as st
import pandas as pd
import requests
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from io import BytesIO

st.set_page_config(page_title="AtheraCLOUD CMC Control Tower", layout="wide")

# --- 사이드바 타임라인 설정 ---
st.sidebar.markdown("---")
st.sidebar.subheader("📅 타임라인 설정")
start_date = st.sidebar.date_input("프로젝트 시작일", pd.to_datetime("2026-03-01"))
target_ind = st.sidebar.date_input("IND 신청 목표일", pd.to_datetime("2026-12-31"))

def create_timeline_excel(dataframe, start_dt):
    output = BytesIO()
    workbook = xlsxwriter.Workbook(output)
    sheet = workbook.add_worksheet('CMC_Timeline')
    
    # 스타일 설정
    header_fmt = workbook.add_format({'bold': True, 'bg_color': '#D7E4BC', 'border': 1})
    date_fmt = workbook.add_format({'num_format': 'yyyy-mm-dd', 'border': 1})
    
    # 헤더 작성
    headers = ['Category', 'Method', 'Development (Start)', 'Development (End)', 'Qualification (End)', 'Stability Start']
    for col, text in enumerate(headers):
        sheet.write(0, col, text, header_fmt)
    
    # 데이터 기반 일정 계산 (표준 리드타임 적용)
    for row_num, (idx, row) in enumerate(dataframe.iterrows(), start=1):
        dev_start = start_dt
        dev_end = dev_start + pd.Timedelta(weeks=6) # 표준 6주
        qual_end = dev_end + pd.Timedelta(weeks=4)  # 표준 4주
        
        sheet.write(row_num, 0, str(row['Method Category']))
        sheet.write(row_num, 1, str(row['Method']))
        sheet.write(row_num, 2, dev_start, date_fmt)
        sheet.write(row_num, 3, dev_end, date_fmt)
        sheet.write(row_num, 4, qual_end, date_fmt)
        
        # 안정성 시험 대상일 경우 표시
        if row['Stability-indicating'] in ['Yes', 'Partial']:
            sheet.write(row_num, 5, qual_end + pd.Timedelta(days=1), date_fmt)
        else:
            sheet.write(row_num, 5, "N/A")
            
    workbook.close()
    return output.getvalue()

# --- 화면 하단 실행 버튼 ---
if st.sidebar.button("📊 마스터 타임라인 추출 (Excel)"):
    excel_data = create_timeline_excel(df, pd.to_datetime(start_date))
    st.sidebar.download_button(
        label="📥 엑셀 파일 다운로드",
        data=excel_data,
        file_name=f"{doc_number}_Timeline_Plan.xlsx",
        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    )

# 1. 클라우드 Secrets 관리 (노션 연동)
try:
    NOTION_TOKEN = st.secrets["NOTION_TOKEN"]
    DATABASE_ID = st.secrets["NOTION_DB_ID"]
except Exception:
    st.error("⚠️ 클라우드 설정(Settings > Secrets)에 NOTION_TOKEN과 NOTION_DB_ID를 입력해주세요.")
    st.stop()

# 2. Notion API 호출 함수
@st.cache_data(ttl=60)
def fetch_notion_data(database_id, token):
    url = f"https://api.notion.com/v1/databases/{database_id}/query"
    headers = {
        "Authorization": f"Bearer {token}",
        "Notion-Version": "2022-06-28",
        "Content-Type": "application/json"
    }
    response = requests.post(url, headers=headers)
    if response.status_code != 200: return pd.DataFrame()
    
    results = response.json().get("results", [])
    data = []
    for page in results:
        props = page.get("properties", {})
        row = {}
        for key, val in props.items():
            p_type = val.get("type")
            if p_type == "title": row[key] = val["title"][0]["plain_text"] if val["title"] else ""
            elif p_type in ["rich_text", "select"]: 
                if p_type == "select": row[key] = val["select"]["name"] if val["select"] else ""
                else: row[key] = val["rich_text"][0]["plain_text"] if val["rich_text"] else ""
            else: row[key] = str(val.get(p_type, ""))
        data.append(row)
    return pd.DataFrame(data)

# --- 메인 UI ---
st.title("🗺️ Tool 1: CMC Master Roadmap (Live Dashboard)")
st.sidebar.header("⚙️ 문서 설정 (Document Setup)")
doc_number = st.sidebar.text_input("문서 번호", value="Athera-CMC-001")

with st.spinner('노션 데이터를 동기화 중입니다...'):
    df = fetch_notion_data(DATABASE_ID, NOTION_TOKEN)

if not df.empty:
    st.success("🟢 노션 데이터베이스 실시간 연동 성공!")
    
    # 컬럼명 유연하게 매핑 (이미지 기반: Method Category 사용)
    # 이미지 캡처본에 'Category'가 비어있으므로 'Method Category'를 대신 사용하도록 설정합니다.
    target_cat_col = "Method Category" if "Method Category" in df.columns else "Category"
    
    # 탭 UI 구현
    cat_list = [c for c in df[target_cat_col].unique() if str(c).strip() and str(c) != 'None']
    
    if cat_list:
        tabs = st.tabs(cat_list)
        for i, cat in enumerate(cat_list):
            with tabs[i]:
                display_df = df[df[target_cat_col] == cat][["Attribute", "Method", "Stability-indicating", "Typical Purpose"]]
                st.dataframe(display_df, use_container_width=True, hide_index=True)
    else:
        st.warning("분류(Category) 데이터가 부족하여 전체 목록을 표시합니다.")
        st.dataframe(df, use_container_width=True)

    # --- CTD Word 생성 로직 통합 ---
    def create_ctd_docx(dataframe, doc_num):
        doc = Document()
        # 폰트 세팅
        style = doc.styles['Normal']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(11)
        style.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
        
        # 타이틀 (국문 크게, 영문 부제목)
        t_kr = doc.add_heading('3.2.S.4 원료의약품의 관리', level=0)
        t_kr.alignment = WD_ALIGN_PARAGRAPH.CENTER
        t_en = doc.add_heading('3.2.S.4 Control of Drug Substance', level=1)
        t_en.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 표 생성
        doc.add_heading('분석 시험법 요약 (Analytical Procedures Summary)', level=2)
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Medium Shading 1 Accent 1'
        hdr = table.rows[0].cells
        hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text = 'CQA', 'Method', 'Stability', 'Purpose'
        
        for _, row in dataframe.iterrows():
            cells = table.add_row().cells
            cells[0].text, cells[1].text = str(row['Attribute']), str(row['Method'])
            cells[2].text, cells[3].text = str(row['Stability-indicating']), str(row['Typical Purpose'])
        
        bio = BytesIO()
        doc.save(bio)
        return bio.getvalue()

    st.markdown("---")
    if st.button("📥 최신 노션 데이터로 CTD Word 추출"):
        word_file = create_ctd_docx(df, doc_number)
        st.download_button("💾 파일 다운로드", word_file, f"{doc_number}_CTD.docx")